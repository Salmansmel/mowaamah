from pdf2docx import Converter
from docx import Document
from deep_translator import GoogleTranslator

def translate_pdf_to_word(pdf_file, word_file):
    # 1. تحويل PDF إلى Word
    cv = Converter(pdf_file)
    cv.convert(word_file)
    cv.close()

    # 2. الترجمة
    doc = Document(word_file)
    translator = GoogleTranslator(source='en', target='ar')

    for p in doc.paragraphs:
        for run in p.runs:
            if run.text.strip():
                try:
                    run.text = translator.translate(run.text)
                except:
                    pass

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            try:
                                run.text = translator.translate(run.text)
                            except:
                                pass

    # 3. الحفظ
    doc.save(word_file)

# التنفيذ
translate_pdf_to_word('How to Use Time-Lapse Photography Function of NVR.pdf', 'arabic_file.docx')